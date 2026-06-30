import docx
from pathlib import Path
from typing import List, Dict, Any

class DocxEmitter:
    def __init__(self, output_dir: Path):
        self.output_dir = output_dir
        
    def emit(self, records: List[Dict[str, Any]]):
        docx_dir = self.output_dir / "resume" / "docx"
        
        for r in records:
            doc = docx.Document()
            
            name_str = r.get("full_name") or f"{r.get('first_name', '')} {r.get('last_name', '')}".strip()
            # Clean filename
            safe_name = name_str.replace(" ", "_").replace("/", "").replace("\\", "")
            filename = f"{safe_name}_CV_{r['uuid'][:4]}.docx"
            
            doc.add_heading(name_str, 0)
            
            # Contact
            contact_parts = []
            if r.get("email"): contact_parts.append(r["email"])
            if r.get("phone"): contact_parts.append(r["phone"])
            if r.get("github"): contact_parts.append(r["github"])
            if r.get("linkedin"): contact_parts.append(r["linkedin"])
            
            doc.add_paragraph(" | ".join(contact_parts))
            
            # Experience
            doc.add_heading("Experience", level=1)
            for exp in r.get("experience", []):
                p = doc.add_paragraph()
                p.add_run(f"{exp.get('title')} at {exp.get('company')}").bold = True
                p.add_run(f" ({exp.get('start_date', '')} to {exp.get('end_date', '')})")
                
            # Education
            doc.add_heading("Education", level=1)
            for edu in r.get("education", []):
                doc.add_paragraph(f"{edu.get('degree')} - {edu.get('institution')}, {edu.get('end_year')}")
                
            # Skills
            doc.add_heading("Skills", level=1)
            doc.add_paragraph(", ".join(r.get("skills", [])))
            
            doc.save(docx_dir / filename)
            r["docx_filename"] = filename
